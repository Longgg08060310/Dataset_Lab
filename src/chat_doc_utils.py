import os
import json
from datetime import datetime
import uuid
import oracledb
import pandas as pd
import traceback
import cv2
import re
import asyncio
import time
import requests
import torch

from paddleocr import PPStructure,draw_structure_result,save_structure_res
from pdf2image import convert_from_path
from langchain.document_loaders import PyPDFLoader, UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
# from SignalR_connection import *
from webscoket_connect import websocket_client
from io import BytesIO
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from PIL import Image
import io
import base64
from bs4 import BeautifulSoup

import warnings
warnings.filterwarnings("ignore")



# 获取当前文件的绝对路径
current_file_path = os.path.abspath(__file__)

# 获取当前文件的根目录
root_directory = os.path.dirname(current_file_path)

with open(os.path.join(root_directory, 'configs', 'password.txt'), 'r', encoding='utf-8') as f:
    password = f.read()


def user_behavior_monitoring(token, module_list):
    now_date = str(datetime.now().date())
    try:
        url = "https://szhlinvma72.apac.bosch.com:53235/api/User/GetLoginUser"
        user_info = requests.get(url=url, headers={'Authorization': token}, verify=False)
        status_code = user_info.status_code
        if status_code == 200:
            user_info = json.loads(user_info.text)
        else:
            return 401
    except:
        return 401

    user_id = user_info['data']['ntAccount']
    department = user_info['data']['displayName']
    oracle_conn = oracledb.connect(user="KSSE_DA", password=password,
                                   dsn="10.8.214.89:1521/KSSEDA.WORLD")
    oracle_cursor = oracle_conn.cursor()
    for module in module_list:
        sql_script = "SELECT * FROM IDOL_USER_BEHAVIOR_MONITORING WHERE DATETIME='{}' and USER_ID='{}' AND MODULE='{}'".format(
            now_date, user_id, module)
        oracle_cursor.execute(sql_script)
        cols = [i[0] for i in oracle_cursor.description]
        cursor_res = oracle_cursor.fetchall()
        if len(cursor_res) == 0:
            sql_script = "INSERT INTO IDOL_USER_BEHAVIOR_MONITORING VALUES(:1, :2, :3, :4, :5)"
            oracle_cursor.execute(sql_script, [user_id, department, module, 1, now_date])  # 执行-插入sql语句
            oracle_conn.commit()
        else:
            # 如果是'CHATKPI' 开头的其他module，为了避免在kpi tree 的接口会疯狂积累次数，所以只取第一次
            if module[:7] != 'CHATKPI' or module == 'CHATKPI_chat':
                count_df = pd.DataFrame(cursor_res, columns=cols)
                usage_count = count_df['COUNT'].values[0] + 1

                sql_query = "UPDATE IDOL_USER_BEHAVIOR_MONITORING SET COUNT = {} \
                                            WHERE DATETIME='{}' and USER_ID='{}' AND MODULE='{}'".format(
                    usage_count, now_date, user_id, module)
                oracle_cursor.execute(sql_query)
                oracle_conn.commit()
    oracle_cursor.close()
    oracle_conn.close()
    return "Done"


class ChatDoc:
    def knowledge_base_list(self, token):
        """
        key user登陆知识库管理页面，看到所有名下的知识库
        """
        url = "https://szhlinvma72.apac.bosch.com:53235/api/User/GetLoginUser"
        user_info = requests.get(url=url, headers={'Authorization': token}, verify=False)
        status_code = user_info.status_code
        if status_code == 401:
            return 401
        elif status_code == 200:
            user_info = json.loads(user_info.text)
            nt_account = user_info['data']['ntAccount']

        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()

        query = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_LIST WHERE INSTR(CONTRIBUTOR, :search_string) > 0"
        platform_cursor.execute(query, {"search_string": nt_account})
        cols = [i[0] for i in platform_cursor.description]
        knowledge_base_res = platform_cursor.fetchall()
        result = []
        if len(knowledge_base_res) == 0:
            return result
        knowledge_base_df = pd.DataFrame(knowledge_base_res, columns=cols)
        knowledge_base_id_list = list(knowledge_base_df['KNOWLEDGEBASE_ID'].values)

        sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST WHERE KNOWLEDGEBASE_ID in {}".format('('+str(knowledge_base_id_list)[1:-1]+')')
        platform_cursor.execute(sql_script)
        cols = [i[0] for i in platform_cursor.description]
        doc_res = platform_cursor.fetchall()
        if len(doc_res) == 0:
            doc_df = pd.DataFrame({'KNOWLEDGEBASE_ID': [-1]})
        else:
            doc_df = pd.DataFrame(doc_res, columns=cols)


        for i in range(len(knowledge_base_df)):
            res_dict = {'knowledge_base_id': str(knowledge_base_df.loc[i, 'KNOWLEDGEBASE_ID']),
                        'knowledge_base_name': knowledge_base_df.loc[i, 'KNOWLEDGEBASE_NAME'],
                        'knowledge_base_description': knowledge_base_df.loc[i, 'KNOWLEDGEBASE_DESCRIPTION'],
                        'doc_count': len(doc_df[doc_df['KNOWLEDGEBASE_ID'] == knowledge_base_df.loc[i, 'KNOWLEDGEBASE_ID']]),
                        'create_time': knowledge_base_df.loc[i, 'CREATE_TIME'],
                        'creator_name': knowledge_base_df.loc[i, 'CREATOR_NAME'],
                        'creator_dept': knowledge_base_df.loc[i, 'CREATOR_DEPT'],
                        'creator_nt': knowledge_base_df.loc[i, 'CREATOR_NT']}
            result.append(res_dict)
        return result

    def knowledge_base_info(self, token, knowledge_base_id):
        url = "https://szhlinvma72.apac.bosch.com:53235/api/User/GetLoginUser"
        user_info = requests.get(url=url, headers={'Authorization': token}, verify=False)
        status_code = user_info.status_code
        if status_code == 401:
            return 401
        elif status_code == 200:
            user_info = json.loads(user_info.text)
            nt_account = user_info['data']['ntAccount']

        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()

        query = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_LIST WHERE KNOWLEDGEBASE_ID = {}".format(knowledge_base_id)
        platform_cursor.execute(query)
        cols = [i[0] for i in platform_cursor.description]
        knowledge_base_df = pd.DataFrame(platform_cursor.fetchall(), columns=cols)

        sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST WHERE KNOWLEDGEBASE_ID = {}".format(knowledge_base_id)
        platform_cursor.execute(sql_script)
        cols = [i[0] for i in platform_cursor.description]
        doc_res = platform_cursor.fetchall()
        if len(doc_res) == 0:
            doc_df = pd.DataFrame({'KNOWLEDGEBASE_ID': [-1]})
        else:
            doc_df = pd.DataFrame(doc_res, columns=cols)

        res_dict = {'knowledge_base_id': knowledge_base_id,
                    'knowledge_base_name': knowledge_base_df.loc[0, 'KNOWLEDGEBASE_NAME'],
                    'knowledge_base_description': knowledge_base_df.loc[0, 'KNOWLEDGEBASE_DESCRIPTION'],
                    'doc_count': len(doc_df[doc_df['KNOWLEDGEBASE_ID'] == knowledge_base_id]),
                    'create_time': knowledge_base_df.loc[0, 'CREATE_TIME'],
                    'creator_name': knowledge_base_df.loc[0, 'CREATOR_NAME'],
                    'creator_dept': knowledge_base_df.loc[0, 'CREATOR_DEPT'],
                    'creator_nt': knowledge_base_df.loc[0, 'CREATOR_NT'],
                    'if_creator': True if knowledge_base_df.loc[0, 'CREATOR_NT'] == nt_account else False}
        return res_dict

    def create_knowledge_base(self, token, knowledge_name, description):
        """
        新建知识库，输入name 和description 后存入表1
        """
        url = "https://szhlinvma72.apac.bosch.com:53235/api/User/GetLoginUser"
        user_info = requests.get(url=url, headers={'Authorization': token}, verify=False)
        user_info = json.loads(user_info.text)
        creator_nt = user_info['data']['ntAccount']
        creator_name = user_info['data']['lastName'] + ' ' + user_info['data']['firstName']
        creator_dept = user_info['data']['displayName']

        # creator_nt = 'NT123'
        # creator_name = 'User123'

        now_time = str(datetime.now())
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "SELECT COUNT(*) FROM IDOL_CHATDOC_KNOWLEDGEBASE_LIST"
        platform_cursor.execute(sql_script)
        count_res = platform_cursor.fetchall()[0][0]

        sql_script = "INSERT INTO IDOL_CHATDOC_KNOWLEDGEBASE_LIST VALUES(:1, :2, :3, :4, :5, :6, :7, :8)"
        contributor_info = [{"contributor_nt": creator_nt, "contributor_name": creator_name,
                             "contributor_dept": creator_dept}]
        platform_cursor.execute(sql_script, [json.dumps(contributor_info), count_res, knowledge_name, description, now_time,
                                             creator_nt, creator_name, creator_dept])  # 执行-插入sql语句
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()

    def edit_knowledge_base(self, knowledge_base_id, knowledge_name, description):
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()

        sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_LIST SET KNOWLEDGEBASE_NAME = '{}', KNOWLEDGEBASE_DESCRIPTION = '{}' \
                    WHERE KNOWLEDGEBASE_ID = {}".format(knowledge_base_id, knowledge_name, description)
        platform_cursor.execute(sql_query)
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()

    def contributor_list(self, knowledge_base_id):
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "SELECT CONTRIBUTOR FROM IDOL_CHATDOC_KNOWLEDGEBASE_LIST WHERE KNOWLEDGEBASE_ID = {}".format(knowledge_base_id)
        platform_cursor.execute(sql_script)
        # print(platform_cursor.fetchall()[0][0])
        contributor_res =[{'contributor_nt': d['contributor_nt'], 'contributor_name': d['contributor_name'], 'contributor_dept': d['contributor_dept']} for d in json.loads(platform_cursor.fetchall()[0][0])]
        return contributor_res

    def add_contributor(self, knowledge_base_id, nt, first_name, last_name, dept):
        """
        向知识库中新增可使用人员
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "SELECT CONTRIBUTOR FROM IDOL_CHATDOC_KNOWLEDGEBASE_LIST WHERE KNOWLEDGEBASE_ID = {}".format(
            knowledge_base_id)
        platform_cursor.execute(sql_script)
        contributor_res = json.loads(platform_cursor.fetchall()[0][0])
        contributor_res.append({"contributor_nt": nt.upper(), "contributor_name": last_name+' '+first_name,
                                "contributor_dept": dept})

        sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_LIST SET CONTRIBUTOR = '{}' WHERE KNOWLEDGEBASE_ID = {}".format(
            json.dumps(contributor_res), knowledge_base_id
        )
        platform_cursor.execute(sql_query)
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()

    def remove_contributor(self, knowledge_base_id, nt):
        """
        向知识库中删除可使用人员
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "SELECT CONTRIBUTOR FROM IDOL_CHATDOC_KNOWLEDGEBASE_LIST WHERE KNOWLEDGEBASE_ID = {}".format(
            knowledge_base_id)
        platform_cursor.execute(sql_script)
        contributor_res = json.loads(platform_cursor.fetchall()[0][0])
        for i in range(len(contributor_res)):
            if contributor_res[i]['contributor_nt'] == nt:
                contributor_res.pop(i)
                break

        sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_LIST SET CONTRIBUTOR = '{}' WHERE KNOWLEDGEBASE_ID = {}".format(
            json.dumps(contributor_res), knowledge_base_id
        )
        platform_cursor.execute(sql_query)
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()

    def upload_file(self, knowledge_base_id, path):
        """
        向知识库中新增文件
        """
        now_time = str(datetime.now())
        doc_id = str(uuid.uuid4())
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "INSERT INTO IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST VALUES(:1, :2, :3, :4, :5)"
        platform_cursor.execute(sql_script, [int(knowledge_base_id), doc_id, path, now_time, 0])  # 执行-插入sql语句
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()

    def doc_list(self, knowledge_base_id):
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST WHERE KNOWLEDGEBASE_ID = {}".format(knowledge_base_id)
        platform_cursor.execute(sql_script)
        cols = [i[0] for i in platform_cursor.description]
        result = []
        doc_res = platform_cursor.fetchall()
        if len(doc_res) == 0:
            return result
        doc_df = pd.DataFrame(doc_res, columns=cols)

        sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT WHERE KNOWLEDGEBASE_ID = {}".format(knowledge_base_id)
        platform_cursor.execute(sql_script)
        cols = [i[0] for i in platform_cursor.description]
        chunking_res = platform_cursor.fetchall()
        if len(chunking_res) == 0:
            chunking_df = pd.DataFrame({'DOC_ID': [0], 'CHUNK_ID': [0]})
        else:
            chunking_df = pd.DataFrame(chunking_res, columns=cols)

        for i in range(len(doc_df)):
            file_name = doc_df.loc[i, 'DOC_ADDRESS'].split('/')[-1]
            file_type = 'pdf' if file_name.split('.')[-1].lower() == 'pdf' else 'ppt' if file_name.split('.')[
                        -1].lower() in ['ppt', 'pptx'] else 'doc' if file_name.split('.')[-1].lower() in [
                        'doc', 'docx'] else 'excel' if file_name.split('.')[-1].lower() in ['xls', 'xlsx'] else ''
            if doc_df.loc[i, 'STATUS'] == '1':
                status = 'Done'
                chunk_number = len(chunking_df[chunking_df['DOC_ID'] == doc_df.loc[i, 'DOC_ID']]['CHUNK_ID'].unique())
            elif doc_df.loc[i, 'STATUS'] == '0':
                status = 'New'
                chunk_number = None
            else:
                status = 'Failed'
                chunk_number = None
            result.append({'doc_id': doc_df.loc[i, 'DOC_ID'], 'doc_name': file_name, 'doc_type': file_type,
                           'status': status, 'chunk_number': chunk_number, 'upload_time': doc_df.loc[i, 'UPLOAD_TIME']})
        platform_cursor.close()
        platform_conn.close()
        return result

    def remove_file(self, doc_id):
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "delete from IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST where DOC_ID='{}'".format(doc_id)
        platform_cursor.execute(sql_script)
        platform_conn.commit()

        sql_script = "delete from IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT where DOC_ID='{}'".format(doc_id)
        platform_cursor.execute(sql_script)
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()

    def chunk_method_list(self, doc_id):
        """
        根据文件的类型判断chunk method
        """
        return ['Document']


    ##############################################################
    #                      PDF Chunking                          #
    ##############################################################
    @staticmethod
    def _parsing_pdf(image_path, page):
        if not os.path.exists(image_path):
            print(f"文件 {image_path} 不存在！")
        table_engine = PPStructure(show_log=True, image_orientation=True, lang='ch')
        img = cv2.imread(image_path)
        pps_result = table_engine(img)

        with open(os.path.join(root_directory, 'output', 'page_content.txt'), 'w', encoding='utf-8') as f:
            f.write(page.page_content)

        result = []
        for type_dict in pps_result:
            if type_dict['type'] in ['title', 'text', 'header', 'figure']:
                region = []
                content = []
                for content_dict in type_dict['res']:
                    content.append(content_dict['text'])
                    region.append(content_dict['text_region'])
                # content_text = ''.join(content)
                type_res = {'type': 'text', 'region': region, 'content': content}
                result.append(type_res)
            elif type_dict['type'] == 'table':
                content = type_dict['res']['html']
                prompt = f"""
                    我在对文档进行RAG实现对文档的问答，现在想对表格内容进行处理，我会给你一段html格式的内容，请帮我转化为一段自然语言来描述整个表格的内容, 
                    形式如下：第1行xxx列的值是xxx；...（请自动识别表头对上面的内容进行替换）,以便后续进行embedding。如果出现缺失表头的情况，请根据你的经验适当补充。
                    请直接给我转换过后的内容,不需要多余的其他内容！html内容为：{content}
                    """
                conversation = [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]

                response = azure_client.chat.completions.create(
                    model="gpt-4o",  # 使用 GPT-4 模型
                    messages=conversation,
                    temperature=0.5,  # 设置生成的温度
                    stream=False  # 启用流式输出
                )
                response = response.choices[0].message.content

                type_res = {'type': 'table', 'region': type_dict['bbox'], 'content': [response]}
                result.append(type_res)
        return result

    @staticmethod
    def _cal_region_pdf(position_matrix):
        first_locat = []
        second_locat = []
        for position_list in position_matrix:
            for position in position_list:
                first_locat.append(position[0])
                second_locat.append(position[1])
        if len(first_locat) > 0:
            return [[min(first_locat), min(second_locat)], [max(first_locat), max(second_locat)]]
        else:
            return [[None, None], [None, None]]

    @staticmethod
    def _text_split(text):
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=300)
        chunks = text_splitter.split_text(text)
        return chunks

    def _chunking_pdf(self, parsing_result):
        chunk_result = []
        for result_dict in parsing_result:
            # if result_dict['type'] == 'text':
            content_text = '.'.join(result_dict['content'])
            if len(content_text) < 2000:
                if result_dict['type'] == 'text':
                    result_dict['region'] = self._cal_region_pdf(result_dict['region'])
                else:
                    result_dict['region'] = [[result_dict['region'][0], result_dict['region'][1]],
                                             [result_dict['region'][2], result_dict['region'][3]]]
                result_dict['content'] = content_text
                chunk_result.append(result_dict)
            else:
                chunks = self._text_split(content_text)
                for chunk in chunks:
                    start_idx, end_idx = None, None
                    for i in range(len(result_dict['content'])):
                        if result_dict['content'][i] in chunk and start_idx is None:
                            start_idx = i
                        elif result_dict['content'][i] not in chunk and start_idx is not None:
                            end_idx = i - 1
                            break
                    if start_idx is not None and end_idx is None:
                        end_idx = len(result_dict['content']) - 1
                    try:
                        region = self._cal_region_pdf(result_dict['region'][start_idx: end_idx + 1])
                        chunk_result.append({'type': result_dict['type'], 'region': region, 'content': chunk})
                    except:
                        if result_dict['type'] == 'text':
                            region = self._cal_region_pdf(result_dict['region'])
                        else:
                            region = [[result_dict['region'][0], result_dict['region'][1]],
                                      [result_dict['region'][2], result_dict['region'][3]]]
                        chunk_result.append({'type': result_dict['type'], 'region': region, 'content': chunk})
            # elif result_dict['type'] == 'table':
            #     # 表格怎么chunking可以解决embedding 和 前端显示的矛盾
            #     # 怎么解决表格被切分到两页的问题
            #     # 表格的表头识别不准确的问题
            #     content_df = result_dict['content']
            #     content_df = content_df.astype(str)
            #     content_df.fillna('', inplace=True)
            #     chunk_content = {'columns': content_df.columns, 'rows': []}
            #     for i in range(len(content_df)):
            #         chunk_content['rows'].append(dict(content_df.loc[i]))
            #
            #     chunk_result.append({'type': 'table', 'region': result_dict['region'], 'content': chunk_content})
        return chunk_result


    ##############################################################
    #                      WORD Chunking                         #
    ##############################################################
    @staticmethod
    def _get_html_from_word(doc_path: str, doc_name) -> str:
        response = requests.get(doc_path, verify=False)

        # 将文件内容加载到内存
        file_stream = BytesIO(response.content)
        doc = Document(file_stream)
        html_content = []
        image_count = 0  # 图片计数器

        # 遍历文档中的所有元素
        for element in doc.element.body:
            # 判断是段落
            if isinstance(element, CT_P):
                para = next((p for p in doc.paragraphs if p._element == element), None)
                if para:
                    style_name = para.style.name if para.style else "Normal"
                    text = para.text.strip()
                    for run in para.runs:
                        drawing_element = run._element.find(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                        if drawing_element is not None:
                            rel_id = drawing_element.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            if rel_id:
                                rel = doc.part.rels.get(rel_id)
                                if rel and "image" in rel.target_ref:
                                    image_data = rel.target_part.blob
                                    image_filename = f"{doc_name}_image_{image_count}.png"
                                    image_file = os.path.join(root_directory, 'output', image_filename)
                                    with open(image_file, 'wb') as file:
                                        file.write(image_data)
                                    with open(image_file, "rb") as image:
                                        files = {"file": (image_filename, image, "image/png")}  # 定义文件信息

                                        upload_url = 'https://szhlinvma75.apac.bosch.com:59108/api/upload'
                                        response = requests.post(upload_url, files=files, verify=False)
                                        image_path = response.json()['data'][0]['url']
                                        print(image_path)
                                    # image_path = os.path.join(output_dir, image_filename)
                                    # Image.open(io.BytesIO(image_data)).save(image_path, format="PNG")
                                    html_content.append(f"<img src='{image_path}' alt='Image {image_count}' />")
                                    image_count += 1
                    if text:
                        if style_name.startswith("Heading 1"):
                            html_content.append(f"<h1>{text}</h1>")
                        elif style_name.startswith("Heading 2"):
                            html_content.append(f"<h2>{text}</h2>")
                        elif style_name.startswith("Heading 3"):
                            html_content.append(f"<h3>{text}</h3>")
                        else:
                            html_content.append(f"<p>{text}</p>")

            # 判断是表格
            elif isinstance(element, CT_Tbl):
                table = next((tbl for tbl in doc.tables if tbl._element == element), None)
                if table:
                    html_content.append("<table border='1'>")
                    for row in table.rows:
                        html_content.append("<tr>")
                        for cell in row.cells:
                            cell_html_content = []
                            # 遍历单元格内容（可能包含文本和图片）
                            for paragraph in cell.paragraphs:
                                # 提取单元格中的文本
                                cell_text = paragraph.text.strip()
                                if cell_text:
                                    cell_html_content.append(cell_text)

                                # 提取单元格中的图片
                                for run in paragraph.runs:
                                    drawing_element = run._element.find(
                                        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                                    if drawing_element is not None:
                                        rel_id = drawing_element.get(
                                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                                        if rel_id:
                                            rel = doc.part.rels.get(rel_id)
                                            if rel and "image" in rel.target_ref:
                                                image_data = rel.target_part.blob
                                                image_filename = f"{doc_name}_image_{image_count}.png"
                                                image_file = os.path.join(root_directory, 'output', image_filename)
                                                with open(image_file, 'wb') as file:
                                                    file.write(image_data)
                                                with open(image_file, "rb") as image:
                                                    files = {"file": (image_filename, image, "image/png")}  # 定义文件信息

                                                    upload_url = 'https://szhlinvma75.apac.bosch.com:59108/api/upload'
                                                    response = requests.post(upload_url, files=files, verify=False)
                                                    image_path = response.json()['data'][0]['url']
                                                    print(image_path)

                                                # image_path = os.path.join(output_dir, image_filename)
                                                # Image.open(io.BytesIO(image_data)).save(image_path, format="PNG")
                                                # 将图片插入到单元格 HTML 内容中
                                                cell_html_content.append(
                                                    f"<img src='{image_path}' alt='Image {image_count}' />")
                                                image_count += 1

                            # 将单元格内容添加到表格的 HTML
                            html_content.append(f"<td>{' '.join(cell_html_content)}</td>")
                        html_content.append("</tr>")
                    html_content.append("</table>")

        return "\n".join(html_content)

    @staticmethod
    def _table2NL(html_data):
        # Parse HTML content
        soup = BeautifulSoup(html_data, 'html.parser')
        rows = soup.find_all('tr')

        # Determine if the first row is a header by checking if all cells are identical
        first_row_cells = [cell.get_text(strip=True) for cell in rows[0].find_all('td')]
        if len(set(first_row_cells)) == 1:
            # Use the second row as header
            header = [cell.get_text(strip=True) for cell in rows[1].find_all('td')]
            data_start_row = 2
        else:
            # Use the first row as header
            header = first_row_cells
            data_start_row = 1

        # Create natural language descriptions
        descriptions = []
        for i, row in enumerate(rows[data_start_row:], start=1):
            cells = [cell.get_text(strip=True) for cell in row.find_all('td')]
            if any(cells):  # Skip empty rows
                row_description = f"第{i}行中" + ", ".join([f"{header[j]}的值是{cells[j]}" for j in range(len(cells))])
                descriptions.append(row_description)

        # Output result
        result = "\n".join(descriptions)
        return result

    @staticmethod
    def _remove_content_between_a_and_b(text, start_sign, end_sign):
        # 找到第一次出现的 'a' 和最后一次出现的 'b' 的索引
        start = text.find(start_sign)
        end = text.rfind(end_sign)

        # 如果找到的 'a' 在 'b' 之前，则删除两者之间的内容
        if start != -1 and end != -1 and start < end:
            res = text[:start] + text[end:]
        else:
            res = text
        return res

    def _chunking_word(self, html_output):
        soup = BeautifulSoup(html_output, 'html.parser')
        # 查找所有的 h1 标签
        h1_tags = soup.find_all('h1')
        # 存储分块的内容
        chunks = []

        # 获取所有的元素（标签和文本）
        elements = list(soup.descendants)
        # 处理第一个 <h1> 标签之前的内容
        before_first_h1 = ''.join(str(element) for element in elements[:elements.index(h1_tags[0])])
        if before_first_h1:
            # 使用正则表达式查找并分割字符串，保留 <table> 标签
            table_chunks = re.split(r'(<table.*?</table>)', before_first_h1, flags=re.DOTALL)
            # 输出分割后的结果
            for i, table_chunk in enumerate(table_chunks):
                chunks.append(table_chunk.strip())

        # 遍历所有 h1 标签并根据其顺序提取内容
        for i in range(len(h1_tags)):
            start_h1_tag = h1_tags[i]
            h1_content = str(start_h1_tag)  # 包含当前 h1 标签
            # 提取从 start_tag 到 end_tag 之间的所有内容
            h2_tags = start_h1_tag.find_all_next(['h2'])
            h2_content = ''
            h3_tags = start_h1_tag.find_all_next(['h3'])
            h3_content = ''
            chunk = [h1_content]

            # 如果不是最后一个 h1，找到下一个 h1 标签作为结束位置
            if i + 1 < len(h1_tags):
                end_h1_tag = h1_tags[i + 1]
                for element in start_h1_tag.find_all_next():
                    if element == end_h1_tag:
                        chunk_text = ''.join(chunk)
                        chunks.append(chunk_text)
                        break
                    elif element in h2_tags:
                        h2_content = str(element)
                        if len(chunk) > 1:
                            chunk_text = ''.join(chunk)
                            chunks.append(chunk_text)
                            chunk = [h1_content, h2_content]
                        else:
                            chunk.append(str(element))
                    elif element in h3_tags:
                        h3_content = str(element)
                        if len(chunk) > 2:
                            chunk_text = ''.join(chunk)
                            chunks.append(chunk_text)
                            chunk = [h1_content, h2_content, h3_content]
                        else:
                            chunk.append(str(element))
                    else:
                        chunk.append(str(element))

            else:
                # 最后的一个 h1 标签处理，直到文档末尾
                for element in start_h1_tag.find_all_next():
                    if element in h2_tags:
                        h2_content = str(element)
                        if len(chunk) > 1:
                            chunk_text = ''.join(chunk)
                            chunks.append(chunk_text)
                            chunk = [h1_content, h2_content]
                        else:
                            chunk.append(str(element))
                    elif element in h3_tags:
                        h3_content = str(element)
                        if len(chunk) > 2:
                            chunk_text = ''.join(chunk)
                            chunks.append(chunk_text)
                            chunk = [h1_content, h2_content, h3_content]
                        else:
                            chunk.append(str(element))
                    else:
                        chunk.append(str(element))

                chunk_text = ''.join(chunk)
                chunks.append(chunk_text)

        for i in range(len(chunks)):
            # 匹配 <tr> 和 </tr> 中间的内容，包括 <tr> 和 </tr>
            pattern = re.compile(r"(<table.*?</table>)", re.DOTALL)
            matches = pattern.findall(chunks[i])
            chunks[i] = self._remove_content_between_a_and_b(chunks[i], '<table', '</table>')
            chunks[i] = self._remove_content_between_a_and_b(chunks[i], '<tr>', '</td>')
            res = ''
            for match in matches:
                res += self._table2NL(match)
            chunks[i] = chunks[i].replace('</table>', res)
            chunks[i] = chunks[i].replace('</td>', '')
        return chunks

    @staticmethod
    def _embedding(text):
        # 加载模型
        model_name = "/app/multilingual-e5-large-instruct"
        model = SentenceTransformer(model_name, trust_remote_code=True)

        # model_name = "intfloat/multilingual-e5-large-instruct"
        # model = SentenceTransformer(model_name)
        model = model.to(torch.float16)

        # 生成文本嵌入
        embedding = model.encode(text, batch_size=2)
        return [float(v) for v in embedding]


    def start_parsing(self, task_id, knowledge_base_id, doc_id):
        """
        start_parsing, 实时传递parsing的进度
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        start_datetime = datetime.now()
        print(start_datetime)
        if not os.path.exists(os.path.join(root_directory, 'output')):
            os.makedirs(os.path.join(root_directory, 'output'))
        try:
            status_message = {'connectionID': task_id,
                              'category': 'chunking_status', 'from': '', 'to': '',
                              'message':json.dumps({'Process_Begin_At': str(start_datetime),
                                                    'Process_Duration': '0s',
                                                    'Process_Msg': 'Task has been received. Start Parsing...',
                                                    'Status': 0}),
                              'remarks': ''}
            try:
                asyncio.run(websocket_client(status_message))
            except:
                time.sleep(0.5)
                asyncio.run(websocket_client(status_message))

            sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST where DOC_ID = '{}'".format(doc_id)
            platform_cursor.execute(sql_script)
            cols = [i[0] for i in platform_cursor.description]
            doc_address_df = pd.DataFrame(platform_cursor.fetchall(), columns=cols)
            doc_status = doc_address_df['STATUS'].values[0]
            # '-1'表示失败， '0'表示New， '1'表示成功， '2'表示正在跑
            if doc_status == '2':
                status_message = {'connectionID': task_id,
                                  'category': 'chunking_status', 'from': '', 'to': '',
                                  'message': json.dumps({'Status': 2}),
                                  'remarks': ''}
                try:
                    asyncio.run(websocket_client(status_message))
                except:
                    time.sleep(0.5)
                    asyncio.run(websocket_client(status_message))
                return None

            sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST SET STATUS = '{}' WHERE DOC_ID = '{}'".format(2, doc_id)
            platform_cursor.execute(sql_query)
            platform_conn.commit()

            doc_address = doc_address_df['DOC_ADDRESS'].values[0]
            doc_name = '.'.join('_'.join(doc_address.split('/')[-1].split('_')[1:]).split('.')[:-1])

            if doc_address.split('/')[-1].split('.')[-1].lower() == 'pdf':
                response = requests.get(doc_address, verify=False)

                pdf_file = os.path.join(root_directory, 'output', f'{doc_name}.pdf')
                with open(pdf_file, 'wb') as f:
                    f.write(response.content)
                    # 使用 PyPDFLoader 加载 PDF
                    loader = PyPDFLoader(pdf_file)  # 传递文件路径
                    pages = loader.load()

                    # Convert PDF to images
                    # poppler_path = "/usr/bin"
                    images = convert_from_path(pdf_path=pdf_file)
                    # Save each page as a separate image
                    for i, image in enumerate(images):
                        now_datetime = datetime.now()
                        duration = (now_datetime - start_datetime).seconds
                        status_message = {'connectionID': task_id,
                                          'category': 'chunking_status', 'from': '', 'to': '',
                                          'message': json.dumps({'Process_Begin_At': str(start_datetime),
                                                                 'Process_Duration': '%.2fs' % duration,
                                                                 'Process_Msg': 'Current Parsing Status: %d/%d Pages' % (i, len(images)),
                                                                 'Status': 0}),
                                          'remarks': ''}
                        try:
                            asyncio.run(websocket_client(status_message))
                        except:
                            time.sleep(0.5)
                            asyncio.run(websocket_client(status_message))

                        image_file = os.path.join(root_directory, 'output', f"{doc_name}_{i + 1}.png")
                        image.save(image_file, "PNG")
                        with open(image_file, "rb") as image:
                            files = {"file": (f"{doc_name}_{i + 1}.png", image, "image/png")}  # 定义文件信息
                            upload_url = 'https://szhlinvma75.apac.bosch.com:59108/api/upload'
                            response = requests.post(upload_url, files=files, verify=False)
                            image_path = response.json()['data'][0]['url']
                            print(image_path)

                        page = pages[i]
                        parsing_result = self._parsing_pdf(image_file, page)
                        chunk_result = self._chunking_pdf(parsing_result)
                        for chunk_dict in chunk_result:
                            chunk_id = str(uuid.uuid4())
                            embedding_res_list = []
                            embedding_res = self._embedding(chunk_dict['content'])
                            embedding_res_list.append(embedding_res)
                            # if chunk_dict['type'] == 'text':
                            #     embedding_res = self._embedding(chunk_dict['content'])
                            #     embedding_res_list.append(embedding_res)
                            # elif chunk_dict['type'] == 'table':
                            #     for row_dict in chunk_dict['content']['rows']:
                            #         row_text = ''
                            #         for col, value in row_dict.items():
                            #             row_text += '%s的值是%s,' % (col, value)
                            #         embedding_res = self._embedding(row_text)
                            #         embedding_res_list.append(embedding_res)

                            sql_script = "INSERT INTO IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT VALUES(:1, :2, :3, :4, :5, :6, :7, :8)"
                            try:
                                platform_cursor.execute(sql_script, [knowledge_base_id, doc_id, i+1, image_path, chunk_id,
                                                                     json.dumps(chunk_dict, ensure_ascii=False), json.dumps(embedding_res_list, ensure_ascii=False), 'pdf'])  # 执行-插入sql语句
                                platform_conn.commit()
                            except:
                                continue

            elif doc_address.split('/')[-1].split('.')[-1].lower() == 'docx':
                html_output = self._get_html_from_word(doc_address, doc_name)

                now_datetime = datetime.now()
                duration = (now_datetime - start_datetime).seconds
                status_message = {'connectionID': task_id,
                                  'category': 'chunking_status', 'from': '', 'to': '',
                                  'message': json.dumps({'Process_Begin_At': str(start_datetime),
                                                         'Process_Duration': '%.2fs' % duration,
                                                         'Process_Msg': 'Starting Text Split...',
                                                         'Status': 0}),
                                  'remarks': ''}
                try:
                    asyncio.run(websocket_client(status_message))
                except:
                    time.sleep(0.5)
                    asyncio.run(websocket_client(status_message))

                # 将 HTML 写入文件
                txt_file = os.path.join(root_directory, 'output', "%s.txt" % doc_name)
                with open(txt_file, "w", encoding="utf-8") as file:
                    file.write(html_output)
                with open(txt_file, "r", encoding="utf-8") as file:
                    files = {"file": (doc_name+'.txt', file, "text/plain")}  # 定义文件信息
                    upload_url = 'https://szhlinvma75.apac.bosch.com:59108/api/upload'
                    response = requests.post(upload_url, files=files, verify=False)
                    txt_path = response.json()['data'][0]['url']
                    print(txt_path)
                chunks_res_list = self._chunking_word(html_output)
                for i, chunk in enumerate(chunks_res_list):
                    now_datetime = datetime.now()
                    duration = (now_datetime - start_datetime).seconds
                    status_message = {'connectionID': task_id,
                                      'category': 'chunking_status', 'from': '', 'to': '',
                                      'message': json.dumps({'Process_Begin_At': str(start_datetime),
                                                             'Process_Duration': '%.2fs' % duration,
                                                             'Process_Msg': 'Current Parsing Status: %d/%d Chunks' % (i, len(chunks_res_list)),
                                                             'Status': 0}),
                                      'remarks': ''}
                    try:
                        asyncio.run(websocket_client(status_message))
                    except:
                        time.sleep(0.5)
                        asyncio.run(websocket_client(status_message))

                    if chunk.strip():
                        chunk_id = str(uuid.uuid4())
                        embedding_res_list = []
                        embedding_res = self._embedding(chunk)
                        embedding_res_list.append(embedding_res)
                        chunk_dict = {"type": "text", "region": "", "content": chunk}
                        sql_script = "INSERT INTO IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT VALUES(:1, :2, :3, :4, :5, :6, :7, :8)"
                        platform_cursor.execute(sql_script, [knowledge_base_id, doc_id, 1, txt_path, chunk_id,
                                                             json.dumps(chunk_dict, ensure_ascii=False), json.dumps(embedding_res_list), 'docx'])
                        platform_conn.commit()
            else:
                now_datetime = datetime.now()
                duration = (now_datetime - start_datetime).seconds
                sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST SET STATUS = '{}' WHERE DOC_ID = '{}'".format(
                    -1, doc_id)
                platform_cursor.execute(sql_query)
                platform_conn.commit()
                status_message = {'connectionID': task_id,
                                  'category': 'chunking_status', 'from': '', 'to': '',
                                  'message': json.dumps({'Process_Begin_At': str(start_datetime),
                                                         'Process_Duration': '%.2fs' % duration,
                                                         'Process_Msg': 'Current Parsing Status: Failed',
                                                         'Status': -1}),
                                  'remarks': ''}
                try:
                    asyncio.run(websocket_client(status_message))
                except:
                    time.sleep(0.5)
                    asyncio.run(websocket_client(status_message))
                return None

            sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST SET STATUS = '{}' WHERE DOC_ID = '{}'".format(1, doc_id)
            platform_cursor.execute(sql_query)
            platform_conn.commit()
            status_message = {'connectionID': task_id,
                              'category': 'chunking_status', 'from': '', 'to': '',
                              'message': json.dumps({'Process_Begin_At': str(start_datetime),
                                                     'Process_Duration': '%.2fs' % duration,
                                                     'Process_Msg': 'Current Parsing Status: Done',
                                                     'Status': 1}),
                              'remarks': ''}
            try:
                asyncio.run(websocket_client(status_message))
            except:
                time.sleep(0.5)
                asyncio.run(websocket_client(status_message))

        except Exception as e:
            print(str(e))
            traceback.print_exc()
            now_datetime = datetime.now()
            duration = (now_datetime - start_datetime).seconds
            sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_DOC_LIST SET STATUS = '{}' WHERE DOC_ID = '{}'".format(-1, doc_id)
            platform_cursor.execute(sql_query)
            platform_conn.commit()
            status_message = {'connectionID': task_id,
                              'category': 'chunking_status', 'from': '', 'to': '',
                              'message': json.dumps({'Process_Begin_At': str(start_datetime),
                                                     'Process_Duration': '%.2fs' % duration,
                                                     'Process_Msg': 'Current Parsing Status: Failed',
                                                     'Status': -1}),
                              'remarks': ''}
            try:
                asyncio.run(websocket_client(status_message))
            except:
                time.sleep(0.5)
                asyncio.run(websocket_client(status_message))

            sql_script = "DELETE FROM IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT WHERE DOC_ID = '{}'".format(doc_id)
            platform_cursor.execute(sql_script)  # 执行-插入sql语句
            platform_conn.commit()

        platform_cursor.close()
        platform_conn.close()


    def show_parsing_detail(self, doc_id):
        """
        parsing完成后，用户点击文件，可以看到detail 信息
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT where DOC_ID = '{}'".format(doc_id)
        platform_cursor.execute(sql_script)
        cols = [i[0] for i in platform_cursor.description]
        doc_detail_df = pd.DataFrame(platform_cursor.fetchall(), columns=cols)

        # 解析 CLOB 到字典
        def parse_clob_to_dict(clob):
            if clob is None:  # 处理空值
                return None
            if isinstance(clob, oracledb.LOB):  # CLOB 类型
                clob = clob.read()  # 读取 CLOB 全部内容
            try:
                return json.loads(clob)  # 解析 JSON
            except json.JSONDecodeError:
                return clob  # 如果解析失败，保持原字符串

        doc_detail_df['CHUNK_CONTENT'] = doc_detail_df['CHUNK_CONTENT'].apply(parse_clob_to_dict)


        if doc_detail_df['DOC_TYPE'].values[0] == 'pdf':
            result = {'doc_type': 'pdf', 'value': []}
            for name, group in doc_detail_df.groupby('DOC_PAGE_ID'):
                group.reset_index(drop=True, inplace=True)
                result_dict = {'page_no': name, 'image_path': group['PAGE_ADDRESS'].values[0],
                               'chunk_value': []}
                for i in range(len(group)):
                    chunk_dict = {'chunk_id': group.loc[i, 'CHUNK_ID'], 'chunk_region': group.loc[i, 'CHUNK_CONTENT']['region'],
                                  'chunk_type': group.loc[i, 'CHUNK_CONTENT']['type'],
                                  'chunk_content': group.loc[i, 'CHUNK_CONTENT']['content']}
                    result_dict['chunk_value'].append(chunk_dict)
                result['value'].append(result_dict)
        elif doc_detail_df['DOC_TYPE'].values[0] == 'docx':
            result = {'doc_type': 'docx'}
            result['value'] = doc_detail_df['PAGE_ADDRESS'].values[0]
        return result


    def modify_parsing_content(self, doc_type, new_content_dict):
        """
        用户可以对parsing 结果进行修改
        if doc_type == 'pdf':
            new_content_dict: {'chunk_id': xxx, 'new_content': {'type': xx, 'region': xx, 'content': xxx}}
        if doc_type == 'docx':
            new_content_dict: {'knowledge_base_id': xxx, 'doc_id': xxx, 'new_file_path': xxx}
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        if doc_type == 'pdf':
            chunk_type = new_content_dict['new_content']['type']
            embedding_res_list = []
            if chunk_type == 'text':
                new_text = new_content_dict['new_content']['content']
                embedding_res = self._embedding(new_text)
                embedding_res_list.append(embedding_res)
                sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT SET CHUNK_CONTENT = '{}', EMBEDDING = '{}' WHERE CHUNK_ID = '{}'".format(
                    json.dumps(new_content_dict['new_content'], ensure_ascii=False), json.dumps(embedding_res_list), new_content_dict['chunk_id']
                )

            elif chunk_type == 'table':
                for row_dict in new_content_dict['new_content']['content']['rows']:
                    row_text = ''
                    for col, value in row_dict.items():
                        row_text += '%s的值是%s,' % (col, value)
                    embedding_res = self._embedding(row_text)
                    embedding_res_list.append(embedding_res)
                sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT SET CHUNK_CONTENT = '{}', EMBEDDING = '{}' WHERE CHUNK_ID = '{}'".format(
                    json.dumps(new_content_dict['new_content'], ensure_ascii=False), json.dumps(embedding_res_list), new_content_dict['chunk_id']
                )
            platform_cursor.execute(sql_query)
            platform_conn.commit()

        elif doc_type == 'docx':
            sql_query = "DELETE FROM IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT WHERE DOC_ID = '{}'".format(
                new_content_dict['doc_id'])
            platform_cursor.execute(sql_query)
            platform_conn.commit()

            with open(new_content_dict['new_file_path'], "r", encoding="utf-8") as file:
                html_output = file.read()
            chunks_res_list = self._chunking_word(html_output)
            for i, chunk in enumerate(chunks_res_list):
                chunk_id = str(uuid.uuid4())
                embedding_res_list = []
                embedding_res = self._embedding(chunk)
                embedding_res_list.append(embedding_res)
                chunk_dict = {"type": "text", "region": "", "content": chunk}
                sql_script = "INSERT INTO IDOL_CHATDOC_KNOWLEDGEBASE_CHUNKING_RESULT VALUES(:1, :2, :3, :4, :5, :6, :7, :8)"
                platform_cursor.execute(sql_script, [new_content_dict['knowledge_base_id'],
                                                     new_content_dict['doc_id'], 1, new_content_dict['new_file_path'],
                                                     chunk_id, json.dumps(chunk_dict, ensure_ascii=False), json.dumps(embedding_res_list), 'docx'])
                platform_conn.commit()

        platform_cursor.close()
        platform_conn.close()


    def create_assistant(self, name, knowledge_base_id, prompt):
        """
        新建assistant 绑定知识库
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()

        now_time = str(datetime.datetime.now())
        assistant_id = str(uuid.uuid4())
        sql_script = "INSERT INTO IDOL_CHATDOC_KNOWLEDGEBASE_ASSISTANT_LIST VALUES(:1, :2, :3, :4, :5, :6)"
        platform_cursor.execute(sql_script, [assistant_id, name, knowledge_base_id, prompt, now_time])  # 执行-插入sql语句
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()


    def assistant_list(self, knowledge_base_id):
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_ASSISTANT_LIST WHERE KNOWLEDGEBASE_ID = {}".format(knowledge_base_id)
        platform_cursor.execute(sql_script)
        cols = [i[0] for i in platform_cursor.description]
        result = []
        assistant_res = platform_cursor.fetchall()
        if len(assistant_res) == 0:
            return result
        assistant_df = pd.DataFrame(assistant_res, columns=cols)

        assistant_id_list = list(assistant_df['ASSISTANT_ID'].values)
        sql_script = "SELECT * FROM IDOL_CHATDOC_KNOWLEDGEBASE_USER_LIST WHERE ASSISTANT_ID IN {}".format(
            '(' + str(assistant_id_list)[1:-1] + ')')
        platform_cursor.execute(sql_script)
        cols = [i[0] for i in platform_cursor.description]
        user_res = platform_cursor.fetchall()
        if len(user_res) == 0:
            user_df = pd.DataFrame({'ASSISTANT_ID': [0]})
        else:
            user_df = pd.DataFrame(user_res, columns=cols)

        for i in range(len(assistant_df)):
            res_dict = {'assistant_id': assistant_df.loc[i, 'ASSISTANT_ID'],
                        'assistant_name': assistant_df.loc[i, 'ASSISTANT_NAME'],
                        'assistant_prompt': assistant_df.loc[i, 'PROMPT'],
                        'assistant_user': []}
            user_sub_df = user_df[user_df['ASSISTANT_ID'] == assistant_df.loc[i, 'ASSISTANT_ID']]
            user_sub_df.reset_index(drop=True, inplace=True)
            for j in range(len(user_sub_df)):
                res_dict['assistant_user'].append({'user_name': user_sub_df.loc[j, 'LAST_NAME'] + ' ' + user_sub_df.loc[j, 'FIRST_NAME'],
                                                   'user_dept': user_sub_df.loc[j, 'DEPT']})
            result.append(res_dict)
        return result


    def edit_assistant(self, assistant_id, assistant_name, assistant_prompt):
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_query = "UPDATE IDOL_CHATDOC_KNOWLEDGEBASE_ASSISTANT_LIST SET ASSISTANT_NAME = '{}', PROMPT = '{}' WHERE ASSISTANT_ID = '{}'".format(
            assistant_name, assistant_prompt, assistant_id
        )
        platform_cursor.execute(sql_query)
        platform_conn.commit()


    def add_user(self, assistant_id, nt, work_no, first_name, last_name, dept):
        """
        向知识库中新增可使用人员
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "INSERT INTO IDOL_CHATDOC_KNOWLEDGEBASE_USER_LIST VALUES(:1, :2, :3, :4, :5, :6)"
        platform_cursor.execute(sql_script, [assistant_id, nt.upper(), work_no, first_name, last_name, dept])  # 执行-插入sql语句
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()


    def remove_user(self, assistant_id, work_no):
        """
        向知识库中删除可使用人员
        """
        platform_conn = oracledb.connect(user="KSSE_DA", password=password,
                                         dsn="10.8.214.89:1521/KSSEDA.WORLD")
        platform_cursor = platform_conn.cursor()
        sql_script = "delete from IDOL_CHATDOC_KNOWLEDGEBASE_USER_LIST where ASSISTANT_ID='{}' AND WORK_NO='{}'".format(
            assistant_id, work_no)
        platform_cursor.execute(sql_script)
        platform_conn.commit()
        platform_cursor.close()
        platform_conn.close()


if __name__ == '__main__':
    chat_doc_object = ChatDoc()
    # chat_doc_object.create_knowledge_base('test1', 'text_description')
    # res = chat_doc_object.upload_file('0', 'https://szhlinvma72.apac.bosch.com:53196/uploads/605db0a5-255e-431d-bc08-58f3c2e60afc_W AE MFG CE - 1000.29 AE-Wuj ITM_ITML Local Concept.docx')
    # chat_doc_object.start_parsing('111', '2', '5c42d4be-591f-42f7-8151-a488f8e86080')
    result = chat_doc_object.show_parsing_detail('5ecd674c-bdc7-4a04-9af8-37ba842c1f35')
    print(json.dumps(result, ensure_ascii=False))